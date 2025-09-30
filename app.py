from flask import Flask, render_template, request, send_file
import requests
from urllib.parse import urlparse
from collections import defaultdict
from io import BytesIO
from docx import Document

app = Flask(__name__)

def get_repo_structure(github_url, branch="main"):
    path_parts = urlparse(github_url).path.strip("/").split("/")
    if len(path_parts) < 2:
        raise ValueError("Invalid GitHub repo URL")
    owner, repo = path_parts[0], path_parts[1]

    api_url = f"https://api.github.com/repos/{owner}/{repo}/git/trees/{branch}?recursive=1"
    response = requests.get(api_url)
    if response.status_code != 200:
        raise Exception(f"Failed to fetch repo structure: {response.status_code}, {response.text}")

    tree = response.json().get("tree", [])
    return [item["path"] for item in tree]

def build_tree(paths):
    tree = lambda: defaultdict(tree)
    root = tree()
    for path in paths:
        parts = path.split("/")
        current = root
        for part in parts:
            current = current[part]
    return root

def format_tree(d, prefix=""):
    lines = []
    items = list(d.keys())
    for i, key in enumerate(items):
        is_last = (i == len(items) - 1)
        connector = "└── " if is_last else "├── "
        line = prefix + connector + key
        lines.append(line)
        if d[key]:
            extension = "    " if is_last else "│   "
            lines += format_tree(d[key], prefix + extension)
    return lines

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        github_url = request.form.get("github_url")
        branch = request.form.get("branch") or "main"
        try:
            structure = get_repo_structure(github_url, branch)
            tree = build_tree(structure)
            tree_lines = format_tree(tree)
            return render_template(
                "tree.html",
                tree_lines=tree_lines,
                repo_url=github_url,
                github_url=github_url,
                branch=branch
            )
        except Exception as e:
            error = str(e)
            return render_template("index.html", error=error)
    return render_template("index.html")


@app.route("/download", methods=["POST"])
def download():
    github_url = request.form.get("github_url")
    branch = request.form.get("branch") or "main"
    try:
        structure = get_repo_structure(github_url, branch)
        tree = build_tree(structure)
        tree_lines = format_tree(tree)
        
        # Create Word document
        doc = Document()
        doc.add_heading("GitHub Repository Structure", 0)
        doc.add_paragraph(github_url)
        for line in tree_lines:
            doc.add_paragraph(line)
        
        # Save to in-memory file
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return send_file(
            file_stream,
            as_attachment=True,
            download_name="repo_structure.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        return f"Error: {e}"


if __name__ == "__main__":
    app.run(debug=True)
